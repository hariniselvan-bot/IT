from docx import Document

doc = Document()

doc.add_heading("Web Development Notes", 0)
doc.add_paragraph("HTML, CSS & JavaScript\nBeginner to Project-Oriented 🌸")

doc.add_page_break()

doc.add_heading("HTML", level=1)
doc.add_paragraph("HTML is used to create the structure of web pages.")

doc.add_page_break()

doc.add_heading("CSS", level=1)
doc.add_paragraph("CSS is used to style web pages.")

doc.add_page_break()

doc.add_heading("JavaScript", level=1)
doc.add_paragraph("JavaScript adds interactivity to websites.")

doc.save("Girly_Pink_HTML_CSS_JS_Notes.docx")
